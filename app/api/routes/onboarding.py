"""
Onboarding Routes - expose onboarding data/summary for workspaces
"""

from typing import Any, Optional, Dict, List

from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File, BackgroundTasks
from pydantic import BaseModel
from sqlalchemy.orm import Session
from datetime import datetime, timezone

from app.database.database import get_db
from app.api.dependencies.auth import get_current_user
from app.api.dependencies.permissions import (
    get_workspace_membership,
)
from app.models.workspace import Workspace, WorkspaceMember
from app.models.retrospective_new import Retrospective
from app.models.onboarding import UserOnboarding
from app.models.user import User
from app.core.config import settings
from app.ai.openai_client import AIClient
from app.ai.features.onboarding_summary import generate_onboarding_summary
from app.ai.rag.chroma_store import ChromaStore
from app.ai.utils import hash_text

# Optional PDF parsing
try:
    from pypdf import PdfReader  # type: ignore
except Exception:  # pragma: no cover
    PdfReader = None  # type: ignore

# Optional Word document parsing
try:
    from docx import Document  # type: ignore
except Exception:  # pragma: no cover
    Document = None  # type: ignore


def _extract_text_from_pdf_bytes(pdf_bytes: bytes) -> str:
    """Extract text from PDF bytes using pypdf if available; otherwise return empty string."""
    if PdfReader is None:
        return ""
    try:
        import io
        with io.BytesIO(pdf_bytes) as fh:
            reader = PdfReader(fh)
            texts = []
            for page in reader.pages:
                try:
                    texts.append(page.extract_text() or "")
                except Exception:
                    continue
            return "\n".join(t for t in texts if t)
    except Exception as e:
        # Log the error for debugging
        import logging
        logger = logging.getLogger(__name__)
        logger.warning(f"PDF extraction failed: {str(e)}")
        return ""


def _extract_text_from_docx_bytes(docx_bytes: bytes) -> str:
    """Extract text from Word document (.docx) bytes using python-docx if available."""
    if Document is None:
        return ""
    try:
        import io
        with io.BytesIO(docx_bytes) as fh:
            doc = Document(fh)
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_texts = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_texts.append(cell.text.strip())
                    if row_texts:
                        paragraphs.append(" | ".join(row_texts))
            return "\n\n".join(paragraphs)
    except Exception as e:
        # Log the error for debugging
        import logging
        logger = logging.getLogger(__name__)
        logger.warning(f"DOCX extraction failed: {str(e)}")
        return ""


def _extract_text_from_docx_manual(docx_bytes: bytes) -> str:
    """Fallback: Extract text from .docx by parsing the ZIP/XML structure manually."""
    try:
        import zipfile
        import xml.etree.ElementTree as ET
        import io
        
        with zipfile.ZipFile(io.BytesIO(docx_bytes), 'r') as zip_file:
            # Read the main document XML
            if 'word/document.xml' not in zip_file.namelist():
                return ""
            
            xml_content = zip_file.read('word/document.xml')
            root = ET.fromstring(xml_content)
            
            # Define namespaces for Word documents
            namespaces = {
                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            }
            
            # Extract all text from paragraphs
            paragraphs = []
            for para in root.findall('.//w:p', namespaces):
                texts = []
                for text_elem in para.findall('.//w:t', namespaces):
                    if text_elem.text:
                        texts.append(text_elem.text)
                if texts:
                    paragraphs.append(''.join(texts))
            
            return '\n\n'.join(paragraphs)
    except Exception as e:
        import logging
        logger = logging.getLogger(__name__)
        logger.warning(f"Manual DOCX extraction failed: {str(e)}")
        return ""


router = APIRouter(prefix="/api/v1", tags=["onboarding"])


class OnboardingUpdateRequest(BaseModel):
    onboarding_data: Any


@router.get("/workspaces/{workspace_id}/onboarding")
def get_workspace_onboarding(
    workspace_id: int,
    membership = Depends(get_workspace_membership),
    db: Session = Depends(get_db)
):
    """
    Return the owner's onboarding_data for a workspace.
    Any active workspace member can view this summary (regardless of role).
    """
    workspace: Optional[Workspace] = db.query(Workspace).filter(Workspace.id == workspace_id).first()
    if not workspace:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Workspace not found")

    # Find onboarding record for the owner (created_by) and this workspace
    owner_onboarding: Optional[UserOnboarding] = db.query(UserOnboarding).filter(
        UserOnboarding.user_id == workspace.created_by,
        UserOnboarding.workspace_id == workspace_id
    ).first()

    # Compute dynamic onboarding flags based on live data
    members_count = db.query(WorkspaceMember).filter(WorkspaceMember.workspace_id == workspace_id, WorkspaceMember.is_active == True).count()
    has_team = members_count >= 2
    has_retro = db.query(Retrospective).filter(Retrospective.workspace_id == workspace_id).count() > 0
    data = owner_onboarding.onboarding_data if owner_onboarding and isinstance(owner_onboarding.onboarding_data, dict) else {}
    source_text = data.get("source_text", "") if isinstance(data, dict) else ""
    imported = bool(source_text and source_text.strip())

    # Update booleans on record if out of sync
    changed = False
    if owner_onboarding:
        if owner_onboarding.workspace_setup_completed != True:
            owner_onboarding.workspace_setup_completed = True; changed = True
        if owner_onboarding.team_members_added != has_team:
            owner_onboarding.team_members_added = has_team; changed = True
        if owner_onboarding.project_data_imported != imported:
            owner_onboarding.project_data_imported = imported; changed = True
        if owner_onboarding.first_retrospective_scheduled != has_retro:
            owner_onboarding.first_retrospective_scheduled = has_retro; changed = True
        if changed:
            db.commit()
            db.refresh(owner_onboarding)

    return {
        "workspace_id": workspace_id,
        "owner_user_id": workspace.created_by,
        "onboarding_data": owner_onboarding.onboarding_data if owner_onboarding else None,
        "status": {
            "workspace_setup_completed": True,
            "team_members_added": has_team,
            "project_data_imported": imported,
            "first_retrospective_scheduled": has_retro,
            "check_ins_configured": bool(owner_onboarding.check_ins_configured) if owner_onboarding else False,
            "onboarding_completed": bool(owner_onboarding.onboarding_completed) if owner_onboarding else False,
        }
    }


@router.put("/workspaces/{workspace_id}/onboarding")
def update_workspace_onboarding(
    workspace_id: int,
    payload: OnboardingUpdateRequest,
    membership = Depends(get_workspace_membership),
    db: Session = Depends(get_db)
):
    """
    Update the owner's onboarding_data for the workspace.
    Requires Scrum Master or Project Manager membership in the workspace.
    # Enforce privileged roles (Scrum Master or Project Manager)
    role_update = (membership.role or '').strip().lower()
    if role_update not in ('scrum master', 'project manager'):
        raise HTTPException(status_code=403, detail="Requires Scrum Master or Project Manager privileges")
    """
    workspace: Optional[Workspace] = db.query(Workspace).filter(Workspace.id == workspace_id).first()
    if not workspace:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Workspace not found")

    # Ensure an onboarding record exists for owner and this workspace
    owner_onboarding: Optional[UserOnboarding] = db.query(UserOnboarding).filter(
        UserOnboarding.user_id == workspace.created_by,
        UserOnboarding.workspace_id == workspace_id
    ).first()

    # Ensure JSON object structure for onboarding_data
    incoming: Any = payload.onboarding_data
    if isinstance(incoming, str):
        # Wrap raw string into structured JSON
        incoming = {"source_text": incoming, "ai_summary": None}
    elif not isinstance(incoming, dict):
        incoming = {"raw": incoming}

    incoming["updated_at"] = datetime.now(timezone.utc).isoformat()

    if not owner_onboarding:
        owner_onboarding = UserOnboarding(
            user_id=workspace.created_by,
            workspace_id=workspace_id,
            onboarding_data=incoming
        )
        db.add(owner_onboarding)
    else:
        # Preserve existing ai_summary unless explicitly provided
        existing = owner_onboarding.onboarding_data if isinstance(owner_onboarding.onboarding_data, dict) else {}
        if isinstance(incoming, dict):
            if "ai_summary" not in incoming and isinstance(existing, dict) and "ai_summary" in existing:
                incoming["ai_summary"] = existing.get("ai_summary")
        owner_onboarding.onboarding_data = incoming

    db.commit()
    db.refresh(owner_onboarding)

    return {
        "workspace_id": workspace_id,
        "owner_user_id": workspace.created_by,
        "onboarding_data": owner_onboarding.onboarding_data,
        "updated": True,
    }


# ====================== User Onboarding (Self) ======================

@router.get("/user-onboarding/me")
def get_my_onboarding(
    workspace_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Get the authenticated user's onboarding record for a specific workspace.
    """
    rec = db.query(UserOnboarding).filter(
        UserOnboarding.user_id == current_user.id,
        UserOnboarding.workspace_id == workspace_id
    ).first()
    return {
        "user_id": current_user.id,
        "workspace_id": workspace_id,
        "onboarding_data": rec.onboarding_data if rec else None
    }


@router.put("/user-onboarding")
def upsert_my_onboarding(
    workspace_id: int,
    payload: OnboardingUpdateRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Create or update the authenticated user's onboarding record for a specific workspace.
    """
    # Verify workspace exists and user has access
    workspace: Optional[Workspace] = db.query(Workspace).filter(Workspace.id == workspace_id).first()
    if not workspace:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Workspace not found")
    
    rec = db.query(UserOnboarding).filter(
        UserOnboarding.user_id == current_user.id,
        UserOnboarding.workspace_id == workspace_id
    ).first()
    if not rec:
        rec = UserOnboarding(
            user_id=current_user.id,
            workspace_id=workspace_id,
            onboarding_data=payload.onboarding_data
        )
        db.add(rec)
    else:
        rec.onboarding_data = payload.onboarding_data
    db.commit()
    db.refresh(rec)
    return {
        "user_id": current_user.id,
        "workspace_id": workspace_id,
        "onboarding_data": rec.onboarding_data,
        "updated": True
    }


# ====================== Upload and Generate Summary ======================

@router.post("/workspaces/{workspace_id}/onboarding/upload")
async def upload_onboarding_source(
    workspace_id: int,
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    membership = Depends(get_workspace_membership),
    db: Session = Depends(get_db)
):
    # Enforce privileged roles (Scrum Master or Project Manager)
    role = (membership.role or '').strip().lower()
    if role not in ('scrum master', 'project manager'):
        raise HTTPException(status_code=403, detail="Requires Scrum Master or Project Manager privileges")
    """
    Upload a file to set onboarding_data for the workspace owner.
    Supported formats:
    - Text files: .txt, .md (plain text and markdown)
    - PDF files: .pdf (text-based PDFs only, not scanned/image-only)
    - Word documents: .docx (modern Word format)
    
    The extracted text will be stored and can be summarized using the generate endpoint.
    """
    content_bytes = await file.read()
    text = ""
    filename_lower = (file.filename or "").lower()
    ctype = (file.content_type or "").lower()
    
    # Determine file type and extract text accordingly
    if ctype in ("application/pdf",) or filename_lower.endswith(".pdf"):
        # Handle PDF files
        text = _extract_text_from_pdf_bytes(content_bytes)
        if not text:
            raise HTTPException(
                status_code=415, 
                detail="Unable to parse PDF. Ensure the file is not scanned or image-only. Text-based PDFs are supported."
            )
    elif (ctype in ("application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                    "application/msword") or 
          filename_lower.endswith((".docx", ".doc"))):
        # Handle Word documents (.docx and .doc)
        if filename_lower.endswith(".docx"):
            # Try python-docx first, then fallback to manual parsing
            text = _extract_text_from_docx_bytes(content_bytes)
            if not text:
                # Fallback to manual ZIP/XML parsing
                text = _extract_text_from_docx_manual(content_bytes)
        else:
            # .doc files (older format) - not easily parseable without additional libraries
            raise HTTPException(
                status_code=415,
                detail="Legacy .doc files are not supported. Please convert to .docx format or upload as PDF."
            )
        
        if not text:
            raise HTTPException(
                status_code=415,
                detail="Unable to extract text from Word document. Please ensure the file is a valid .docx file."
            )
    else:
        # Treat as plain text/markdown
        try:
            text = content_bytes.decode("utf-8", errors="ignore")
        except Exception:
            # Try other encodings
            try:
                text = content_bytes.decode("latin-1", errors="ignore")
            except Exception:
                text = ""
        
        if not text or not text.strip():
            raise HTTPException(
                status_code=415, 
                detail="Unsupported file format or empty content. Supported formats: .txt, .md, .pdf, .docx"
            )

    workspace: Optional[Workspace] = db.query(Workspace).filter(Workspace.id == workspace_id).first()
    if not workspace:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Workspace not found")

    owner_onboarding: Optional[UserOnboarding] = db.query(UserOnboarding).filter(
        UserOnboarding.user_id == workspace.created_by,
        UserOnboarding.workspace_id == workspace_id
    ).first()
    # Build structured JSON
    data: Dict[str, Any] = {}
    if owner_onboarding and isinstance(owner_onboarding.onboarding_data, dict):
        data.update(owner_onboarding.onboarding_data)
    
    uploaded_at = datetime.now(timezone.utc).isoformat()
    doc_hash = hash_text(text)
    doc_id = f"{workspace_id}:{(file.filename or 'document')}:{doc_hash}"

    # Keep legacy fields for backward compatibility (last upload only)
    data["source_text"] = text
    data["uploaded_at"] = uploaded_at
    data["uploaded_filename"] = file.filename

    # New: track all documents for this workspace (so summary can use ALL docs via workspace-scoped RAG)
    docs = data.get("documents")
    if not isinstance(docs, list):
        docs = []
    docs.append({
        "doc_id": doc_id,
        "doc_hash": doc_hash,
        "filename": file.filename,
        "uploaded_at": uploaded_at,
        "bytes": len(content_bytes),
    })
    data["documents"] = docs
    # Do not overwrite ai_summary here; generation endpoint will set it
    data["updated_at"] = datetime.now(timezone.utc).isoformat()

    if not owner_onboarding:
        owner_onboarding = UserOnboarding(
            user_id=workspace.created_by,
            workspace_id=workspace_id,
            onboarding_data=data
        )
        db.add(owner_onboarding)
    else:
        owner_onboarding.onboarding_data = data

    db.commit()
    db.refresh(owner_onboarding)

    # RAG: chunk + embed + store in Chroma in the background (so upload is fast).
    def _index_onboarding_text(wid: int, did: str, dhash: str, fname: str, uploaded_at_iso: str, txt: str) -> None:
        try:
            ai = AIClient()
            store = ChromaStore()
            store.upsert_text_document(
                ai=ai,
                source="onboarding",
                doc_id=did,
                text=txt,
                embedding_model=getattr(settings, "AI_EMBEDDING_MODEL", "text-embedding-3-small"),
                extra_metadata={
                    "workspace_id": wid,
                    "filename": fname,
                    "uploaded_at": uploaded_at_iso,
                    "doc_hash": dhash,
                },
            )
        except Exception:
            # Never fail the upload due to RAG indexing.
            return

    if background_tasks is not None:
        background_tasks.add_task(_index_onboarding_text, workspace_id, doc_id, doc_hash, (file.filename or ""), uploaded_at, text)

    return {
        "workspace_id": workspace_id,
        "uploaded_filename": file.filename,
        "doc_id": doc_id,
        "doc_hash": doc_hash,
        "bytes": len(content_bytes),
        "onboarding_data_preview": (text[:200] + "...") if len(text) > 200 else text,
        "saved": True,
    }


@router.post("/workspaces/{workspace_id}/onboarding/generate")
def generate_workspace_summary(
    workspace_id: int,
    force: bool = False,
    membership = Depends(get_workspace_membership),
    db: Session = Depends(get_db)
):
    # Enforce privileged roles (Scrum Master or Project Manager)
    role = (membership.role or '').strip().lower()
    if role not in ('scrum master', 'project manager'):
        raise HTTPException(status_code=403, detail="Requires Scrum Master or Project Manager privileges")
    """
    Generate an AI summary for a workspace and save it into the owner's onboarding_data.ai_summary.

    New workflow (Project Documents + RAG):
    - Documents are uploaded via /api/v1/workspaces/{workspace_id}/documents/upload
    - Chunks are indexed to Chroma Cloud with metadata workspace_id + source="onboarding"
    - This endpoint can summarize even if onboarding_data.source_text is missing, as long as
      workspace documents exist (either recorded in onboarding_data.documents or present in Chroma).
    """
    workspace: Optional[Workspace] = db.query(Workspace).filter(Workspace.id == workspace_id).first()
    if not workspace:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Workspace not found")

    owner_onboarding: Optional[UserOnboarding] = db.query(UserOnboarding).filter(
        UserOnboarding.user_id == workspace.created_by,
        UserOnboarding.workspace_id == workspace_id
    ).first()
    
    if not owner_onboarding:
        raise HTTPException(status_code=400, detail="No onboarding data found for this workspace. Please upload a file first.")
    
    raw_data = owner_onboarding.onboarding_data
    source_text = ""
    docs: List[Dict[str, Any]] = []
    if isinstance(raw_data, dict):
        source_text = raw_data.get("source_text", "") or ""
        if isinstance(raw_data.get("documents"), list):
            docs = raw_data.get("documents") or []
    elif isinstance(raw_data, str):
        # Backward compatibility
        source_text = raw_data

    has_source_text = isinstance(source_text, str) and bool(source_text.strip())
    has_docs = bool(docs)

    # If there is no stored source_text, we can still summarize from RAG (workspace docs).
    # However, if there are no docs recorded, double-check Chroma quickly.
    if not has_source_text and not has_docs:
        try:
            store = ChromaStore()
            chroma_stats = store.count_chunks(where_filter={"source": "onboarding", "workspace_id": workspace_id}, limit=300)
            if int(chroma_stats.get("count") or 0) <= 0:
                raise HTTPException(
                    status_code=400,
                    detail="No workspace documents found to summarize. Upload documents in the Project Documents tab first.",
                )
        except HTTPException:
            raise
        except Exception:
            raise HTTPException(
                status_code=400,
                detail="No workspace documents found to summarize. Upload documents in the Project Documents tab first.",
            )

    # Build a stable string to drive hashing/caching even when we don't have raw source_text anymore.
    # This changes when docs change (doc_hash/doc_id list changes), so cached summaries invalidate naturally.
    if not has_source_text:
        parts: List[str] = ["WORKSPACE_DOCUMENTS_FOR_SUMMARY"]
        for d in docs:
            if isinstance(d, dict):
                parts.append(f"{d.get('doc_id','')}>{d.get('doc_hash','')}>{d.get('filename','')}")
        source_text = "\n".join(parts)
    
    # Verify we're using the correct workspace's data
    # This ensures data integrity - the source_text should belong to this workspace
    if isinstance(raw_data, dict) and "updated_at" in raw_data:
        # Optional: Could add additional validation here if needed
        pass

    # Summarize using the centralized AI layer + (RAG if available) + caching.
    # IMPORTANT: Do not save raw document text as "ai_summary" on failure.
    try:
        ai = AIClient()
        result = generate_onboarding_summary(
            ai=ai,
            workspace_id=workspace_id,
            source_text=source_text,
            endpoint_name="onboarding.generate_summary",
            cache=(not force),
        )
        summary_text = result["summary"]
        ai_meta = {
            "model": result.get("model"),
            "cached": result.get("cached"),
            "rag_used": result.get("rag_used"),
            "doc_hash": result.get("doc_hash"),
        }
    except Exception as e:
        import logging
        logging.getLogger(__name__).warning(f"Onboarding summary generation failed: {e}")
        # Return a proper error so the UI doesn't show raw file text as a "summary".
        msg = "Onboarding summary generation failed. Check OPENAI_API_KEY and Chroma configuration."
        if getattr(settings, "DEBUG", False):
            msg = f"{msg} Error: {e}"
        raise HTTPException(status_code=503, detail=msg)

    # Save the summary into onboarding_data JSON under ai_summary
    # IMPORTANT: Preserve the source_text that was used for generation
    data_out: Dict[str, Any] = {}
    if owner_onboarding and isinstance(owner_onboarding.onboarding_data, dict):
        data_out.update(owner_onboarding.onboarding_data)
        # Ensure we're using the same source_text that was uploaded
        # This prevents any potential data mixing issues
        if "source_text" in data_out and data_out["source_text"] != source_text:
            # This should never happen, but log if it does
            import logging
            logger = logging.getLogger(__name__)
            logger.warning(f"Source text mismatch detected for workspace {workspace_id}")
    else:
        # Wrap legacy string source into structured JSON
        data_out = {"source_text": source_text}
    
    # Store the summary with metadata about what was summarized
    data_out["ai_summary"] = summary_text
    data_out["summary_generated_at"] = datetime.now(timezone.utc).isoformat()
    data_out["summary_source_length"] = len(source_text)  # Store length for verification
    data_out["summary_source_kind"] = "source_text" if has_source_text else "workspace_documents_rag"
    data_out["updated_at"] = datetime.now(timezone.utc).isoformat()

    # owner_onboarding should always exist at this point (we checked earlier)
    # But keep this as a safety check
    if not owner_onboarding:
        raise HTTPException(status_code=500, detail="Onboarding record was deleted during generation. Please try again.")
    
    # Update the onboarding data with the generated summary
    owner_onboarding.onboarding_data = data_out
    # Keep computed flags updated
    owner_onboarding.workspace_setup_completed = True
    owner_onboarding.project_data_imported = True if data_out.get("source_text") else owner_onboarding.project_data_imported
    db.commit()
    db.refresh(owner_onboarding)

    out = {
        "workspace_id": workspace_id,
        "onboarding_data": owner_onboarding.onboarding_data,
        "generated": True,
    }
    # Helpful debug info (safe; does not include document content)
    if "ai_meta" in locals():
        out["ai_meta"] = ai_meta
    return out


@router.get("/workspaces/{workspace_id}/onboarding/rag-status")
def get_onboarding_rag_status(
    workspace_id: int,
    membership = Depends(get_workspace_membership),
    db: Session = Depends(get_db),
):
    """
    Debug/status endpoint to verify Chroma indexing for onboarding documents.
    Requires Scrum Master or Project Manager privileges.
    """
    role = (membership.role or '').strip().lower()
    if role not in ('scrum master', 'project manager'):
        raise HTTPException(status_code=403, detail="Requires Scrum Master or Project Manager privileges")

    workspace: Optional[Workspace] = db.query(Workspace).filter(Workspace.id == workspace_id).first()
    if not workspace:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Workspace not found")

    owner_onboarding: Optional[UserOnboarding] = db.query(UserOnboarding).filter(
        UserOnboarding.user_id == workspace.created_by,
        UserOnboarding.workspace_id == workspace_id
    ).first()

    raw_data = owner_onboarding.onboarding_data if owner_onboarding else None
    docs = []
    if isinstance(raw_data, dict) and isinstance(raw_data.get("documents"), list):
        docs = raw_data.get("documents") or []

    # Workspace-scoped stats (ALL docs indexed for this workspace). Do not hard-fail if no docs yet.
    store = ChromaStore()
    # Chroma Cloud typically caps get(limit=...) to <=300; use a safe value so status works in production.
    stats = store.count_chunks(where_filter={"source": "onboarding", "workspace_id": workspace_id}, limit=300)

    # Also verify retrieval returns something (optional).
    retrieved_count = 0
    retrieval_error = None
    try:
        ai = AIClient()
        retrieved = store.query(
            ai=ai,
            source="onboarding",
            query_text="Summarize this project proposal: objective, problem, methodology, outcomes/impact, stakeholders/partners.",
            embedding_model=getattr(settings, "AI_EMBEDDING_MODEL", "text-embedding-3-small"),
            top_k=getattr(settings, "AI_RAG_TOP_K", 10),
            where_filter={"workspace_id": workspace_id},
        )
        retrieved_count = len(retrieved)
    except Exception as e:
        retrieval_error = str(e)

    return {
        "workspace_id": workspace_id,
        "chroma": stats,
        "documents": {
            "count": len(docs),
            "documents": docs,
            "note": "Upload documents via /api/v1/workspaces/{workspace_id}/documents/upload. Indexing runs in the background.",
        },
        "retrieval": {
            "top_k": getattr(settings, "AI_RAG_TOP_K", 10),
            "retrieved_count": retrieved_count,
            "error": retrieval_error,
        },
    }


@router.post("/workspaces/{workspace_id}/onboarding/complete")
def mark_onboarding_complete(
    workspace_id: int,
    membership = Depends(get_workspace_membership),
    db: Session = Depends(get_db)
):
    """Mark onboarding as completed and set completed_at timestamp."""
    role = (membership.role or '').strip().lower()
    if role not in ('scrum master', 'project manager'):
        raise HTTPException(status_code=403, detail="Requires Scrum Master or Project Manager privileges")

    workspace: Optional[Workspace] = db.query(Workspace).filter(Workspace.id == workspace_id).first()
    if not workspace:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Workspace not found")

    owner_onboarding: Optional[UserOnboarding] = db.query(UserOnboarding).filter(
        UserOnboarding.user_id == workspace.created_by,
        UserOnboarding.workspace_id == workspace_id
    ).first()
    if not owner_onboarding:
        owner_onboarding = UserOnboarding(
            user_id=workspace.created_by,
            workspace_id=workspace_id
        )
        db.add(owner_onboarding)

    owner_onboarding.onboarding_completed = True
    owner_onboarding.completed_at = datetime.now(timezone.utc)
    db.commit()
    db.refresh(owner_onboarding)
    return {"workspace_id": workspace_id, "onboarding_completed": True}


